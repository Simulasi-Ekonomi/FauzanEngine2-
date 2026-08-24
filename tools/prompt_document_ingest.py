#!/usr/bin/env python3
"""Bounded, non-executing prompt-context ingestion for DOCX, PDF, text, and RAR containers."""

from __future__ import annotations

import argparse
import hashlib
import io
import json
import sys
from dataclasses import asdict, dataclass
from pathlib import Path
from typing import Iterable

import rarfile
from docx import Document
from pypdf import PdfReader


MAX_INPUT_BYTES = 8 * 1024 * 1024
MAX_ARCHIVE_ENTRIES = 128
MAX_ARCHIVE_UNCOMPRESSED_BYTES = 16 * 1024 * 1024
MAX_PDF_PAGES = 256
MAX_DOCUMENT_CHARS = 16 * 1024
MAX_CONTEXT_CHARS = 64 * 1024
ALLOWED_MEMBER_SUFFIXES = {".docx", ".pdf", ".txt", ".md"}


class DocumentIngestError(ValueError):
    """Input could not be safely accepted into a prompt context."""


@dataclass(frozen=True)
class PromptDocument:
    source: str
    kind: str
    sha256: str
    characters: int
    text: str


def _digest(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()


def _bounded_text(text: str) -> str:
    normalized = "\n".join(line.rstrip() for line in text.replace("\x00", "").splitlines()).strip()
    if not normalized:
        raise DocumentIngestError("document_has_no_extractable_text")
    return normalized[:MAX_DOCUMENT_CHARS]


def _read_docx(data: bytes) -> str:
    try:
        document = Document(io.BytesIO(data))
    except Exception as error:  # Parser errors are data errors, never executable instructions.
        raise DocumentIngestError("invalid_docx") from error
    parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            values = [cell.text.strip() for cell in row.cells]
            if any(values):
                parts.append(" | ".join(values))
    return _bounded_text("\n".join(parts))


def _read_pdf(data: bytes) -> str:
    try:
        reader = PdfReader(io.BytesIO(data), strict=True)
        if reader.is_encrypted:
            raise DocumentIngestError("encrypted_pdf_rejected")
        if len(reader.pages) > MAX_PDF_PAGES:
            raise DocumentIngestError("pdf_page_limit_exceeded")
        return _bounded_text("\n".join((page.extract_text() or "") for page in reader.pages))
    except DocumentIngestError:
        raise
    except Exception as error:
        raise DocumentIngestError("invalid_pdf") from error


def _parse_member(source: str, suffix: str, data: bytes) -> PromptDocument:
    if len(data) == 0 or len(data) > MAX_INPUT_BYTES:
        raise DocumentIngestError("member_size_limit_exceeded")
    suffix = suffix.lower()
    if suffix == ".docx":
        text = _read_docx(data)
        kind = "docx"
    elif suffix == ".pdf":
        text = _read_pdf(data)
        kind = "pdf"
    elif suffix in {".txt", ".md"}:
        try:
            text = _bounded_text(data.decode("utf-8"))
        except UnicodeDecodeError as error:
            raise DocumentIngestError("text_must_be_utf8") from error
        kind = suffix[1:]
    else:
        raise DocumentIngestError("unsupported_document_type")
    return PromptDocument(source=source, kind=kind, sha256=_digest(data), characters=len(text), text=text)


def _read_rar(path: Path, data: bytes) -> list[PromptDocument]:
    if len(data) > MAX_INPUT_BYTES:
        raise DocumentIngestError("archive_size_limit_exceeded")
    try:
        with rarfile.RarFile(path) as archive:
            if archive.needs_password():
                raise DocumentIngestError("encrypted_archive_rejected")
            entries = [entry for entry in archive.infolist() if not entry.isdir()]
            if not entries or len(entries) > MAX_ARCHIVE_ENTRIES:
                raise DocumentIngestError("archive_entry_limit_exceeded")
            total = sum(entry.file_size for entry in entries)
            if total > MAX_ARCHIVE_UNCOMPRESSED_BYTES:
                raise DocumentIngestError("archive_uncompressed_limit_exceeded")
            documents: list[PromptDocument] = []
            for entry in entries:
                suffix = Path(entry.filename).suffix.lower()
                if suffix not in ALLOWED_MEMBER_SUFFIXES:
                    continue
                if entry.file_size == 0 or entry.file_size > MAX_INPUT_BYTES:
                    raise DocumentIngestError("archive_member_size_limit_exceeded")
                with archive.open(entry) as member:
                    member_bytes = member.read(MAX_INPUT_BYTES + 1)
                if len(member_bytes) != entry.file_size or len(member_bytes) > MAX_INPUT_BYTES:
                    raise DocumentIngestError("archive_member_read_limit_exceeded")
                documents.append(_parse_member(f"{path.name}:{entry.filename}", suffix, member_bytes))
            if not documents:
                raise DocumentIngestError("archive_has_no_supported_documents")
            return documents
    except DocumentIngestError:
        raise
    except Exception as error:
        raise DocumentIngestError("invalid_rar") from error


def ingest_paths(paths: Iterable[Path]) -> list[PromptDocument]:
    documents: list[PromptDocument] = []
    total_characters = 0
    for path in paths:
        resolved = path.resolve(strict=True)
        if not resolved.is_file():
            raise DocumentIngestError("input_must_be_regular_file")
        data = resolved.read_bytes()
        suffix = resolved.suffix.lower()
        additions = _read_rar(resolved, data) if suffix == ".rar" else [_parse_member(resolved.name, suffix, data)]
        for document in additions:
            total_characters += document.characters
            if total_characters > MAX_CONTEXT_CHARS:
                raise DocumentIngestError("aggregate_prompt_context_limit_exceeded")
            documents.append(document)
    if not documents:
        raise DocumentIngestError("no_input_documents")
    return documents


def main() -> int:
    parser = argparse.ArgumentParser(description="Extract bounded, non-executing document context for a prompt planner.")
    parser.add_argument("inputs", nargs="+", type=Path)
    parser.add_argument("--output", type=Path, required=True)
    args = parser.parse_args()
    try:
        documents = ingest_paths(args.inputs)
        payload = {"version": 1, "documents": [asdict(document) for document in documents]}
        args.output.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
        print(f"PROMPT_DOCUMENT_INGEST_OK documents={len(documents)} chars={sum(document.characters for document in documents)}")
        return 0
    except (OSError, DocumentIngestError) as error:
        print(f"PROMPT_DOCUMENT_INGEST_REJECTED reason={error}", file=sys.stderr)
        return 2


if __name__ == "__main__":
    raise SystemExit(main())
